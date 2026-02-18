from fastapi import HTTPException
from io import BytesIO
import os
import re
import json
from typing import List, Dict,Union
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from openpyxl import Workbook
from openpyxl.styles import Font
from openpyxl.utils import get_column_letter
from src.utils.utils import parse_markdown_table
from svglib.svglib import svg2rlg
from reportlab.graphics import renderPM
from reportlab.lib.enums import TA_CENTER
from src.schemas.schemas import ExportExcelRequest,ExportPdfRequest,ExportPngRequest,ExportWordRequest
from src.db.redis_service import RedisService
from src.config.redis_config import Config
from fastapi.responses import StreamingResponse




class ExportService:
   
    def __init__(self):
        self.redis_service = RedisService(Config.get_redis_config())   
       
    @staticmethod
    def export_pdf(content: str, title: str) -> BytesIO:
        buffer = BytesIO()
        title = title or "Financial Report"

        content = content or ""

        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            leftMargin=36,
            rightMargin=36,
            topMargin=36,
            bottomMargin=36,
        )

        styles = getSampleStyleSheet()
        story = []

        # ---------- TITLE ----------
        title_style = ParagraphStyle(
            "Title",
            parent=styles["Heading1"],
            fontSize=20,
            fontName="Helvetica-Bold",
            textColor=colors.HexColor("#0f172a"),
            alignment=TA_CENTER,
            spaceAfter=24,
        )

        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 12))

        # ---------- CUSTOM STYLES ----------
        h2_style = ParagraphStyle(
            "H2",
            fontSize=14,
            fontName="Helvetica-Bold",
            textColor=colors.HexColor("#0f172a"),
            spaceBefore=16,
            spaceAfter=8,
            borderBottom=1,
            borderColor=colors.HexColor("#3b82f6"),
            paddingBottom=4,
        )

        h3_style = ParagraphStyle(
            "H3",
            fontSize=11,
            fontName="Helvetica-Bold",
            textColor=colors.HexColor("#1e293b"),
            spaceBefore=12,
            spaceAfter=6,
            leftIndent=8,
            borderLeft=3,
            borderColor=colors.HexColor("#3b82f6"),
            paddingLeft=6,
        )

        normal_style = ParagraphStyle(
            "NormalText",
            fontSize=10,
            leading=14,
            textColor=colors.HexColor("#475569"),
        )

        # ---------- PARSE CONTENT ----------
        lines = content.split("\n")
        i = 0

        while i < len(lines):
            line = lines[i].strip()

            # Empty line
            if not line:
                story.append(Spacer(1, 6))
                i += 1
                continue

            # ---------- TABLE ----------
            # ---------- TABLE ----------
            if line.startswith("|"):
                table, next_i = parse_markdown_table(lines, i)
                
                # ✅ Check FIRST before using table[0]
                if not table or not table[0]:
                    i = next_i
                    continue
                
                # ✅ NOW it's safe to use table[0]
                col_widths = [doc.width / len(table[0])] * len(table[0])
                
                wrapped_rows = []
                for r_idx, row in enumerate(table):
                    wrapped_row = []
                    for cell in row:
                        wrapped_row.append(
                            Paragraph(
                                str(cell),
                                ParagraphStyle(
                                    "Cell",
                                    fontSize=9,
                                    leading=11,
                                    textColor=colors.white if r_idx == 0 else colors.black,
                                ),
                            )
                        )
                    wrapped_rows.append(wrapped_row)
                
                # ✅ Table creation OUTSIDE the loop (not indented inside for loop)
                t = Table(wrapped_rows, colWidths=col_widths)
                t.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0f172a")),
                            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("LEFTPADDING", (0, 0), (-1, -1), 6),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                            ("TOPPADDING", (0, 0), (-1, -1), 6),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                        ]
                    )
                )

                story.append(t)
                story.append(Spacer(1, 12))
                i = next_i
                continue
            # ---------- H2 ----------
            if line.startswith("## "):
                story.append(
                    Paragraph(line.replace("## ", ""), h2_style)
                )
                i += 1
                continue

            # ---------- H3 ----------
            if line.startswith("### "):
                story.append(
                    Paragraph(line.replace("### ", ""), h3_style)
                )
                i += 1
                continue

            # ---------- NORMAL TEXT ----------
            clean = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", line)
            story.append(Paragraph(clean, normal_style))
            i += 1

        # ---------- BUILD ----------
        doc.build(story)
        buffer.seek(0)
        return buffer

    
    @staticmethod
    def export_word(content: str, title: str) -> BytesIO:
        document = Document()
        title = title or "Financial Report"

        content = content or ""
        # ---------- TITLE ----------
        title_p = document.add_heading(level=1)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run(title)
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(15, 23, 42)


        # ---------- STYLES ----------
        def add_h2(text):
            p = document.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(15, 23, 42)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)

        def add_h3(text):
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.15)
            run = p.add_run(text)
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(30, 41, 59)

        def add_normal(text):
            p = document.add_paragraph()
            parts = re.split(r"(\*\*.*?\*\*)", text)

            for part in parts:
                run = p.add_run(part.replace("**", ""))
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(71, 85, 105)
                if part.startswith("**"):
                    run.bold = True

        # ---------- CONTENT PARSE ----------
        lines = content.split("\n")
        i = 0

        while i < len(lines):
            line = lines[i].strip()

            if not line:
                i += 1
                continue

            # ---------- TABLE ----------
            if line.startswith("|"):
                table_data, next_i = parse_markdown_table(lines, i)
                if table_data and table_data[0]:
                    rows, cols = len(table_data), len(table_data[0])
                    table = document.add_table(rows=rows, cols=cols)
                    table.style = "Table Grid"

                    for r in range(rows):
                        for c in range(cols):
                            cell = table.cell(r, c)
                            cell.text = table_data[r][c]
                            for run in cell.paragraphs[0].runs:
                                run.font.size = Pt(9)
                                if r == 0:
                                    run.bold = True

                    i = next_i
                    continue

            # ---------- H2 ----------
            if line.startswith("## "):
                add_h2(line.replace("## ", ""))
                i += 1
                continue

            # ---------- H3 ----------
            if line.startswith("### "):
                add_h3(line.replace("### ", ""))
                i += 1
                continue

            # ---------- NORMAL ----------
            add_normal(line)
            i += 1

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer


    # ---------------- EXCEL ----------------

    @staticmethod
    def export_excel(
        columns: List[str],
        rows: List[dict],
        sheet_name: str = "Financial Data"
    ) -> BytesIO:
        
        sheet_name = sheet_name or "Financial Data"

        wb = Workbook()
        ws = wb.active
        ws.title = sheet_name

        # Header
        ws.append(columns)
        for cell in ws[1]:
            cell.font = Font(bold=True)

        # Rows
        for row in rows:
            ws.append([row.get(col, "") for col in columns])

        # Auto width
        for idx, col in enumerate(columns, 1):
            ws.column_dimensions[get_column_letter(idx)].width = min(len(col) + 5, 40)

        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer
    
    
  
    @staticmethod
    def export_png(svg_input: str, width: int = 1920, height: int = 1120) -> BytesIO:
        try:
            svg_string = svg_input.strip()

            if not svg_string.startswith('<svg'):
                raise ValueError("Invalid SVG input")
            
            svg_io = BytesIO(svg_string.encode('utf-8'))
            drawing = svg2rlg(svg_io)

            if not drawing:
                raise ValueError("Failed to parse SVG")

            # Get original dimensions
            original_width = drawing.width
            original_height = drawing.height
            
            if not original_width or not original_height:
                raise ValueError("SVG has invalid dimensions")

            # Calculate scale to fit within target dimensions while maintaining aspect ratio
            scale_x = width / original_width
            scale_y = height / original_height
            scale = min(scale_x, scale_y)

            # Calculate actual output dimensions
            output_width = original_width * scale
            output_height = original_height * scale

            # Apply scaling to the drawing
            drawing.scale(scale, scale)
            
            # Set the drawing dimensions for rendering
            drawing.width = output_width
            drawing.height = output_height

            # Render to PNG with proper dimensions
            png_bytes = renderPM.drawToString(
                drawing,
                fmt="PNG",
                dpi=150,
                bg=0xFFFFFF
            )

            buffer = BytesIO(png_bytes)
            buffer.seek(0)
            return buffer

        except Exception as e:
            raise ValueError(f"Failed to convert SVG to PNG: {str(e)}")
    
            
    def export_pdf_handler(self, TenantId: str, SessionId: str, req: ExportPdfRequest) -> StreamingResponse:
        
        self.redis_service.validate_tenant_session(
            TenantId,
            SessionId
        )

        msg = self.redis_service.get_data_by_index(
            TenantId=TenantId,
            SessionId=SessionId,
            index=req.index
        )
        
        if not msg:
            raise HTTPException(
                status_code=410,  
                detail="Data has expired or is no longer available"
            )

        if msg.get("role") != "assistant":
            raise HTTPException(
                status_code=400, 
                detail="PDF export is only available for assistant messages"
            )

        content = msg.get("content")
        if not content:
            raise HTTPException(
                status_code=400, 
                detail="No content available for export"
            )

        buffer = self.export_pdf(
            content=content,
            title=req.title
        )

        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename=report_{req.index}.pdf"
            }
        )

    def export_word_handler(self, TenantId: str, SessionId: str,req: ExportWordRequest)-> StreamingResponse:
        
        self.redis_service.validate_tenant_session(
            TenantId,
            SessionId
        )
        
        msg = self.redis_service.get_data_by_index(
            TenantId=TenantId,
            SessionId=SessionId,
            index=req.index
        )
        
        if not msg:
            raise HTTPException(
                status_code=410, 
                detail="Data has expired or is no longer available"
            )

        if msg.get("role") != "assistant":
            raise HTTPException(
                status_code=400, 
                detail="Word export is only available for assistant messages"
            )

        content = msg.get("content")
        if not content:
            raise HTTPException(
                status_code=400, 
                detail="No content available for export"
            )

        buffer = self.export_word(
            content=content,
            title=req.title
        )
        
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=report_{req.index}.docx"
            }
        )

    
    def export_excel_handler(self, TenantId: str, SessionId: str, req: ExportExcelRequest) -> StreamingResponse:
        
        self.redis_service.validate_tenant_session(TenantId, SessionId)
        
        msg = self.redis_service.get_data_by_index(
            TenantId=TenantId,
            SessionId=SessionId,
            index=req.index
        )
        
        if not msg:
            raise HTTPException(
                status_code=410,
                detail="Data has expired or is no longer available"
            )
        
        if msg.get("role") != "system":
            raise HTTPException(
                status_code=400,
                detail="Excel export is only available for system messages"
            )
        
        raw_content = msg.get("content")
        
        # Parse JSON content
        try:
            rows = json.loads(raw_content) if isinstance(raw_content, str) else raw_content
        except json.JSONDecodeError:
            raise HTTPException(
                status_code=400,
                detail="Invalid data format - unable to parse content"
            )
        
        
        # ✅ BETTER VERSION:
        if not rows or not isinstance(rows, list):
            raise HTTPException(
                status_code=400,
                detail="No data available for export"
            )

        if not isinstance(rows[0], dict):
            raise HTTPException(
                status_code=400,
                detail="Invalid data structure - expected list of objects"
            )
        
        buffer = self.export_excel(
            columns=list(rows[0].keys()),
            rows=rows,
            sheet_name=req.sheet_name
        )
        
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={
                "Content-Disposition": f"attachment; filename=data_{req.index}.xlsx"
            }
        )
    
    
    def export_png_handler(self, TenantId: str, SessionId: str, req: ExportPngRequest) -> StreamingResponse:
        
        self.redis_service.validate_tenant_session(TenantId, SessionId)
        
        msg = self.redis_service.get_data_by_index(
            TenantId=TenantId,
            SessionId=SessionId,
            index=req.index
        )
        
        if not msg:
            raise HTTPException(
                status_code=410,
                detail="Data has expired or is no longer available"
            )
        
        if msg.get("role") != "assistant":
            raise HTTPException(
                status_code=400,
                detail="PNG export is only available for assistant messages"
            )
        
        svg = msg.get("content", "").strip()
        
        if not svg:
            raise HTTPException(
                status_code=400,
                detail="No content available for export"
            )
        
        if not svg.startswith("<svg"):
            raise HTTPException(
                status_code=400,
                detail="Invalid content format - expected SVG data"
            )
        
        try:
            buffer = self.export_png(
                svg_input=svg,
                width=req.width,
                height=req.height
            )
        except ValueError as e:
            raise HTTPException(status_code=400, detail=str(e))
            
        return StreamingResponse(
            buffer,
            media_type="image/png",
            headers={
                "Content-Disposition": f"attachment; filename=chart_{req.index}.png"
            }
        )
     
    def generate_pdf_from_text(self,
        content: str,
        title: str,
        output_dir: str,
        file_name: str
    ) -> Dict[str, Union[bool,str]]:

        buffer = self.export_pdf(
            content=content,
            title=title
        )

        os.makedirs(output_dir, exist_ok=True)

        file_path = os.path.join(output_dir, file_name)

        with open(file_path, "wb") as f:
            f.write(buffer.read())

        return {
            "success": True,
            "file_path": file_path
        }
